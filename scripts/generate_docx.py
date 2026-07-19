#!/usr/bin/env python3
"""Generate a tailored Word resume or cover letter from YAML data.

Usage:
    generate_docx.py resume <data.yaml> <output.docx>
    generate_docx.py cover-letter <data.yaml> <output.docx>

The YAML schemas are documented in profile/profile.yaml (resume) and the
cover letter example in applications/_example/. This script only formats
data it's given — all tailoring/wording decisions happen upstream in the
Claude Code session, not here.
"""
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
TEXT = RGBColor(0x22, 0x22, 0x22)


def set_margins(doc, inches=0.6):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


def build_resume(data: dict, out_path: Path):
    doc = Document()
    set_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    contact = data["contact"]

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(contact["name"].upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    contact_line = " | ".join(
        v for v in [
            contact.get("location"),
            contact.get("phone"),
            contact.get("email"),
            contact.get("linkedin"),
        ] if v
    )
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_p.add_run(contact_line)
    run.font.size = Pt(10)

    if data.get("headline"):
        headline_p = doc.add_paragraph()
        headline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = headline_p.add_run(data["headline"])
        run.italic = True
        run.font.size = Pt(10.5)

    if data.get("summary"):
        add_heading(doc, "Professional Summary")
        for para in data["summary"].strip().split("\n\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(para.strip())
            run.font.size = Pt(10.5)

    if data.get("core_competencies"):
        add_heading(doc, "Core Competencies")
        for group in data["core_competencies"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(group["category"] + ": ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = p.add_run(", ".join(group["items"]))
            run2.font.size = Pt(10.5)

    if data.get("professional_strengths"):
        add_heading(doc, "Professional Strengths")
        for s in data["professional_strengths"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(s["name"] + ": ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = p.add_run(s["detail"])
            run2.font.size = Pt(10.5)

    if data.get("experience"):
        add_heading(doc, "Professional Experience")
        for job in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{job['title']} – {job['company']}")
            run.bold = True
            run.font.size = Pt(11)
            run_dates = p.add_run(f"    {job['dates']}")
            run_dates.italic = True
            run_dates.font.size = Pt(10)
            add_bullets(doc, job["bullets"])

    if data.get("tools_and_technologies"):
        add_heading(doc, "Tools & Technologies")
        for category, items in data["tools_and_technologies"].items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(category + ": ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = p.add_run(", ".join(items))
            run2.font.size = Pt(10.5)

    if data.get("education"):
        add_heading(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            text = edu["degree"]
            if edu.get("school"):
                text += f" – {edu['school']}"
            run = p.add_run(text)
            run.font.size = Pt(10.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_cover_letter(data: dict, out_path: Path):
    doc = Document()
    set_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    contact = data["contact"]

    name_p = doc.add_paragraph()
    run = name_p.add_run(contact["name"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT

    contact_line = " | ".join(
        v for v in [
            contact.get("location"),
            contact.get("phone"),
            contact.get("email"),
            contact.get("linkedin"),
        ] if v
    )
    contact_p = doc.add_paragraph()
    run = contact_p.add_run(contact_line)
    run.font.size = Pt(10)

    doc.add_paragraph()

    if data.get("date"):
        doc.add_paragraph(data["date"])

    if data.get("recipient"):
        doc.add_paragraph(data["recipient"])

    doc.add_paragraph()

    doc.add_paragraph(data.get("salutation", "Dear Hiring Manager,"))

    for para in data["body_paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(para.strip())
        run.font.size = Pt(11)

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(contact["name"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    if len(sys.argv) != 4 or sys.argv[1] not in ("resume", "cover-letter"):
        print(__doc__)
        sys.exit(1)

    kind, data_path, out_path = sys.argv[1], Path(sys.argv[2]), Path(sys.argv[3])
    data = yaml.safe_load(data_path.read_text())

    if kind == "resume":
        build_resume(data, out_path)
    else:
        build_cover_letter(data, out_path)

    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
